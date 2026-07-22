#!/usr/bin/env python3
"""Legacy generator for public/demo/stress-styles.docx.

The shipped fixture is the curated QA Word (DOCX Converter Stress Test v1.0).
Do not regenerate unless you intentionally replace that sample.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_LINE_SPACING, WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "public" / "demo" / "stress-styles.docx"
ASSETS = ROOT / "public" / "demo" / "_stress_assets"


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_run_font(run, name: str, size_pt: float | None = None, bold: bool | None = None, italic: bool | None = None, color: RGBColor | None = None) -> None:
    """Bind fonts that exist in the Alpine demo image (Liberation / DejaVu / Noto)."""
    run.font.name = name
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:cs"), name)
    # Prefer Noto CJK for east-asian runs when body font is Liberation/DejaVu.
    east = "Noto Sans CJK JP" if "Mono" not in name and "Serif" not in name else name
    if "Serif" in name:
        east = "Noto Serif CJK JP"
    r_fonts.set(qn("w:eastAsia"), east)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _paragraph_border(paragraph, color: str = "1B4F72", size: str = "18") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "8")
        el.set(qn("w:color"), color)
        p_bdr.append(el)
    p_pr.append(p_bdr)


def _paragraph_shading(paragraph, hex_color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)


def _set_section_columns(section, num: int = 2, space_twips: int = 360) -> None:
    """Enable multi-column layout on a section (OOXML w:cols)."""
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")


def _make_images() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}

    logo = ASSETS / "logo.png"
    img = Image.new("RGB", (640, 160), (12, 74, 110))
    draw = ImageDraw.Draw(img)
    draw.rectangle((20, 20, 620, 140), outline=(255, 255, 255), width=4)
    draw.ellipse((40, 40, 140, 120), fill=(232, 90, 48))
    draw.polygon([(180, 120), (240, 40), (300, 120)], fill=(46, 204, 113))
    draw.rectangle((340, 40, 440, 120), fill=(241, 196, 15))
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 36)
        font_sm = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 18)
    except OSError:
        font = ImageFont.load_default()
        font_sm = font
    draw.text((470, 55), "NOWO", fill=(255, 255, 255), font=font)
    draw.text((470, 100), "Word → PDF stress", fill=(200, 220, 235), font=font_sm)
    img.save(logo)
    paths["logo"] = logo

    chart = ASSETS / "chart.png"
    img = Image.new("RGB", (720, 360), (248, 250, 252))
    draw = ImageDraw.Draw(img)
    draw.rectangle((40, 30, 680, 320), outline=(45, 55, 72), width=2)
    bars = [(80, 280, 140, 80, (37, 99, 235)), (180, 280, 240, 140, (16, 185, 129)), (280, 280, 340, 60, (245, 158, 11)), (380, 280, 440, 200, (239, 68, 68)), (480, 280, 540, 110, (139, 92, 246)), (580, 280, 640, 240, (6, 182, 212))]
    for x0, y0, x1, y1, color in bars:
        draw.rectangle((x0, y1, x1, y0), fill=color)
    draw.line((60, 280, 660, 280), fill=(45, 55, 72), width=2)
    draw.text((50, 10), "Quarterly conversion throughput (synthetic)", fill=(30, 41, 59), font=font_sm)
    img.save(chart)
    paths["chart"] = chart

    photo = ASSETS / "photo.png"
    img = Image.new("RGB", (400, 280), (0, 0, 0))
    draw = ImageDraw.Draw(img)
    for y in range(280):
        r = int(20 + y * 0.6)
        g = int(60 + (280 - y) * 0.4)
        b = int(120 + y * 0.2)
        draw.line([(0, y), (400, y)], fill=(r % 256, g % 256, b % 256))
    draw.ellipse((120, 60, 280, 220), outline=(255, 255, 255), width=6)
    draw.text((90, 240), "Embedded raster stress", fill=(255, 255, 255), font=font_sm)
    img.save(photo)
    paths["photo"] = photo

    return paths


def build() -> Path:
    images = _make_images()
    doc = Document()

    # --- Styles ---
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Liberation Sans"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    for level, size, color in ((1, 28, RGBColor(0x0C, 0x4A, 0x6E)), (2, 20, RGBColor(0x15, 0x65, 0xC0)), (3, 14, RGBColor(0x1E, 0x3A, 0x5F))):
        style = styles[f"Heading {level}"]
        style.font.color.rgb = color
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.name = "Liberation Sans"

    if "Callout" not in [s.name for s in styles]:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
        callout.font.name = "Liberation Sans"
        callout.font.size = Pt(10)
        callout.font.italic = True
        callout.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    # --- Section margins / header-footer ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("WordToPdfBundle · stress-styles.docx  ·  confidential sample")
    _set_run_font(run, "Liberation Sans", 8, color=RGBColor(0x64, 0x74, 0x8B))

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Page ")
    _set_run_font(run, "Liberation Sans", 9, color=RGBColor(0x47, 0x55, 0x69))
    _add_page_number(fp)
    run = fp.add_run("  ·  LibreOffice Writer export fidelity probe")
    _set_run_font(run, "Liberation Sans", 9, color=RGBColor(0x47, 0x55, 0x69))

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = first_footer.paragraphs[0].add_run("COVER — not for production data")
    _set_run_font(run, "Liberation Sans", 8, italic=True, color=RGBColor(0x94, 0xA3, 0xB8))

    # ========== COVER ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(str(images["logo"]), width=Inches(5.8))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(28)
    run = title.add_run("STRESS STYLES PACK")
    _set_run_font(run, "Liberation Serif", 32, bold=True, color=RGBColor(0x0C, 0x4A, 0x6E))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Push LibreOffice Writer → PDF to the edge")
    _set_run_font(run, "Liberation Sans", 16, italic=True, color=RGBColor(0xE8, 0x5A, 0x30))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(18)
    run = meta.add_run(
        "Headings · character runs · tables · images · lists · hyperlinks ·\n"
        "headers/footers · page breaks · landscape section · unicode · code"
    )
    _set_run_font(run, "Liberation Sans", 11, color=RGBColor(0x33, 0x41, 0x55))

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner.paragraph_format.space_before = Pt(36)
    _paragraph_shading(banner, "0C4A6E")
    run = banner.add_run("  nowo-tech/word-to-pdf-bundle  ·  demo/symfony8  ")
    _set_run_font(run, "DejaVu Sans Mono", 12, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    doc.add_page_break()

    # ========== TOC-ish overview ==========
    doc.add_heading("1. Document map", level=1)
    intro = doc.add_paragraph()
    run = intro.add_run("This sample is intentionally loud. ")
    _set_run_font(run, "Liberation Sans", 11)
    run = intro.add_run("It is not a pretty marketing PDF")
    _set_run_font(run, "Liberation Sans", 11, bold=True, color=RGBColor(0xB9, 0x1C, 0x1C))
    run = intro.add_run(" — it is a layout gauntlet for ")
    _set_run_font(run, "Liberation Sans", 11)
    run = intro.add_run("soffice --headless --convert-to pdf")
    _set_run_font(run, "DejaVu Sans Mono", 10, color=RGBColor(0x0F, 0x76, 0x6E))
    run = intro.add_run(".")
    _set_run_font(run, "Liberation Sans", 11)

    call = doc.add_paragraph(style="Callout")
    _paragraph_border(call, "E85A30", "24")
    _paragraph_shading(call, "FFF7ED")
    run = call.add_run("Callout box: ")
    _set_run_font(run, "Liberation Sans", 10, bold=True, color=RGBColor(0xC2, 0x41, 0x0C))
    run = call.add_run(
        "Expect imperfect round-trips on exotic Word features. "
        "Judge success by readable hierarchy, table grid integrity, and embedded images."
    )
    _set_run_font(run, "Liberation Sans", 10, italic=True, color=RGBColor(0x7C, 0x2D, 0x12))

    for i, item in enumerate(
        [
            "Character formatting zoo (fonts, colors, highlight, strike, super/sub)",
            "Nested numbered + bullet lists",
            "Complex tables (merged cells, zebra, numeric alignment)",
            "Raster images (logo, chart, photo) with captions",
            "Hyperlinks and mixed scripts (Latin, Ελληνικά, 中文, العربية, emoji)",
            "Landscape section with wide financial table",
        ],
        start=1,
    ):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        _set_run_font(run, "Liberation Sans", 11)

    # ========== Character zoo ==========
    doc.add_heading("2. Character formatting zoo", level=1)
    doc.add_heading("2.1 Mixed runs in one paragraph", level=2)
    zoo = doc.add_paragraph()
    parts = [
        ("Normal Liberation Sans. ", dict(name="Liberation Sans", size=11)),
        ("Bold. ", dict(name="Liberation Sans", size=11, bold=True)),
        ("Italic. ", dict(name="Liberation Sans", size=11, italic=True)),
        ("Underline. ", dict(name="Liberation Sans", size=11, underline=True)),
        ("Double underline. ", dict(name="Liberation Sans", size=11, underline=WD_UNDERLINE.DOUBLE)),
        ("Strike. ", dict(name="Liberation Sans", size=11, strike=True)),
        ("Red danger. ", dict(name="Liberation Sans", size=11, bold=True, color=RGBColor(0xDC, 0x26, 0x26))),
        ("Blue link-like. ", dict(name="Liberation Sans", size=11, color=RGBColor(0x25, 0x63, 0xEB), underline=True)),
        ("Liberation Serif. ", dict(name="Liberation Serif", size=12, italic=True)),
        ("Liberation Mono. ", dict(name="Liberation Mono", size=10)),
        ("H₂O + E=mc² ", dict(name="Liberation Sans", size=11)),
        ("ALL CAPS SMALL. ", dict(name="Liberation Sans", size=9, bold=True, color=RGBColor(0x33, 0x41, 0x55))),
    ]
    for text, opts in parts:
        run = zoo.add_run(text)
        _set_run_font(
            run,
            opts["name"],
            opts.get("size"),
            bold=opts.get("bold"),
            italic=opts.get("italic"),
            color=opts.get("color"),
        )
        if opts.get("underline") is True:
            run.underline = True
        elif opts.get("underline"):
            run.underline = opts["underline"]
        if opts.get("strike"):
            run.font.strike = True

    # Explicit super/sub
    chem = doc.add_paragraph()
    run = chem.add_run("Chemistry & physics: H")
    _set_run_font(run, "Liberation Sans", 12)
    run = chem.add_run("2")
    run.font.subscript = True
    _set_run_font(run, "Liberation Sans", 12)
    run = chem.add_run("O · CO")
    _set_run_font(run, "Liberation Sans", 12)
    run = chem.add_run("2")
    run.font.subscript = True
    _set_run_font(run, "Liberation Sans", 12)
    run = chem.add_run(" · Einstein E=mc")
    _set_run_font(run, "Liberation Sans", 12)
    run = chem.add_run("2")
    run.font.superscript = True
    _set_run_font(run, "Liberation Sans", 12)

    hi = doc.add_paragraph()
    run = hi.add_run("Highlight sweep: ")
    _set_run_font(run, "Liberation Sans", 11, bold=True)
    for label, color in (
        ("yellow", WD_COLOR_INDEX.YELLOW),
        ("bright green", WD_COLOR_INDEX.BRIGHT_GREEN),
        ("turquoise", WD_COLOR_INDEX.TURQUOISE),
        ("pink", WD_COLOR_INDEX.PINK),
    ):
        run = hi.add_run(f" [{label}] ")
        run.font.highlight_color = color
        _set_run_font(run, "Liberation Sans", 11)

    # Alignments
    doc.add_heading("2.2 Paragraph alignment & spacing", level=2)
    for align, label in (
        (WD_ALIGN_PARAGRAPH.LEFT, "Left-aligned body copy with a slightly longer line so wrapping is visible across the page margins."),
        (WD_ALIGN_PARAGRAPH.CENTER, "Centered line — often used for titles inside sections."),
        (WD_ALIGN_PARAGRAPH.RIGHT, "Right-aligned meta line · 2026-07-22 · v1.0 stress"),
        (WD_ALIGN_PARAGRAPH.JUSTIFY, "Justified paragraph: the quick brown fox jumps over the lazy dog. " * 4),
    ):
        p = doc.add_paragraph(label)
        p.alignment = align
        p.paragraph_format.space_after = Pt(10)

    indent = doc.add_paragraph("First-line indent + hanging feel for a quotation-like block that LibreOffice must reflow carefully.")
    indent.paragraph_format.first_line_indent = Cm(1.25)
    indent.paragraph_format.left_indent = Cm(0.5)

    # ========== Lists ==========
    doc.add_heading("3. Lists (nested)", level=1)
    doc.add_heading("3.1 Bullets", level=2)
    bullets = [
        (0, "Policy intake"),
        (1, "KYC documents"),
        (2, "Passport / DNI scan"),
        (2, "Proof of address"),
        (1, "Risk questionnaire"),
        (0, "Quotation pipeline"),
        (1, "Word template fill (WordTemplateBundle)"),
        (1, "HTML → DOCX (HtmlToWordBundle)"),
        (1, "DOCX → PDF (this bundle)"),
    ]
    for level, text in bullets:
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.75 * level)

    doc.add_heading("3.2 Numbered procedure", level=2)
    for step in (
        "Upload .docx via demo form or hit the stress sample route.",
        "RuntimeRequirementsChecker asserts soffice is present.",
        "LibreOfficeProcessRunner isolates a temp workspace.",
        "PdfExporter streams BinaryFileResponse to the browser.",
    ):
        doc.add_paragraph(step, style="List Number")

    # ========== Tables ==========
    doc.add_heading("4. Tables under pressure", level=1)
    doc.add_heading("4.1 Product matrix (zebra + header shade)", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["SKU", "Plan", "Premium €", "Tax %", "Status"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, "0C4A6E")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                _set_run_font(run, "Liberation Sans", 10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    rows = [
        ("NOW-HOME-01", "Hogar Esencial", "18.90", "21", "Active"),
        ("NOW-HOME-02", "Hogar Plus", "27.40", "21", "Active"),
        ("NOW-AUTO-11", "Auto Terceros", "14.10", "21", "Paused"),
        ("NOW-LIFE-07", "Vida Temporal", "9.75", "0", "Active"),
        ("NOW-PET-03", "Mascotas", "6.20", "21", "Trial"),
    ]
    for r_i, row_data in enumerate(rows):
        row = table.add_row()
        for c_i, value in enumerate(row_data):
            cell = row.cells[c_i]
            cell.text = value
            if r_i % 2 == 1:
                _set_cell_shading(cell, "E0F2FE")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if c_i in (2, 3) else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    _set_run_font(run, "Liberation Sans", 10)

    doc.add_heading("4.2 Merged cells (schedule)", level=2)
    merge_table = doc.add_table(rows=4, cols=4)
    merge_table.style = "Table Grid"
    merge_table.cell(0, 0).merge(merge_table.cell(0, 3))
    merge_table.cell(0, 0).text = "Conversion SLA calendar — merged header"
    _set_cell_shading(merge_table.cell(0, 0), "E85A30")
    for paragraph in merge_table.cell(0, 0).paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            _set_run_font(run, "Liberation Sans", 11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    labels = [
        ("Mon–Thu", "09:00–18:00", "P50 < 3s", "P99 < 12s"),
        ("Friday", "09:00–15:00", "P50 < 4s", "P99 < 15s"),
        ("Weekend", "On-call only", "Best effort", "No SLA"),
    ]
    for r, values in enumerate(labels, start=1):
        for c, value in enumerate(values):
            merge_table.cell(r, c).text = value

    # ========== Images ==========
    doc.add_heading("5. Embedded images", level=1)
    doc.add_paragraph("Chart-like PNG (synthetic bars):")
    doc.add_picture(str(images["chart"]), width=Inches(5.9))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run("Figure 1 — Synthetic throughput chart (embedded PNG)")
    _set_run_font(run, "Liberation Sans", 9, italic=True, color=RGBColor(0x64, 0x74, 0x8B))

    doc.add_paragraph("Gradient photo with overlay circle:")
    doc.add_picture(str(images["photo"]), width=Inches(3.6))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ========== Hyperlinks & unicode ==========
    doc.add_heading("6. Hyperlinks & international text", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Repository: ")
    _set_run_font(run, "Liberation Sans", 11)

    def add_hyperlink_run(paragraph, text: str, url: str) -> None:
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(color)
        r_pr.append(u)
        new_run.append(r_pr)
        text_el = OxmlElement("w:t")
        text_el.text = text
        new_run.append(text_el)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    add_hyperlink_run(p, "github.com/nowo-tech/WordToPdfBundle", "https://github.com/nowo-tech/WordToPdfBundle")

    multi = doc.add_paragraph()
    run = multi.add_run(
        "Español · English · Français · Deutsch · Ελληνικά · "
        "中文（简体） · 日本語 · العربية · עברית · "
        "Emoji stress: 🐘📄✨🔥  ·  Symbols: € £ ¥ ₩  © ® ™  § ¶"
    )
    _set_run_font(run, "Liberation Sans", 11)

    # ========== Code block ==========
    doc.add_heading("7. Monospace “code” block", level=1)
    code = doc.add_paragraph()
    _paragraph_shading(code, "0F172A")
    _paragraph_border(code, "334155", "12")
    run = code.add_run(
        "$ php bin/console nowo:word-to-pdf:check\n"
        "$ curl -OJ http://localhost:8022/stress.pdf\n"
        "# LibreOffice Writer must be installed on the host/container"
    )
    _set_run_font(run, "DejaVu Sans Mono", 9, color=RGBColor(0xE2, 0xE8, 0xF0))

    # ========== Two-column section ==========
    doc.add_page_break()
    cols_section = doc.add_section()
    cols_section.orientation = WD_ORIENT.PORTRAIT
    cols_section.page_width = Cm(21.0)
    cols_section.page_height = Cm(29.7)
    cols_section.top_margin = Cm(2.0)
    cols_section.bottom_margin = Cm(2.0)
    cols_section.left_margin = Cm(1.8)
    cols_section.right_margin = Cm(1.8)
    _set_section_columns(cols_section, num=2, space_twips=400)

    # Heading spanning: first para before column flow still uses section cols in Word/LO
    doc.add_heading("7b. Two-column newspaper layout", level=1)
    left_bits = [
        "Column A — intake narrative. When a broker uploads a dense Word pack, LibreOffice must reflow continuous text across balanced columns without clipping glyphs at the gutter.",
        "Keep an eye on hyphenation, justification, and whether the column gap (gutter) survives PDF export. This is a classic fidelity stress for Writer → PDF.",
        "Lorem-style filler keeps density high: policies, endorsements, schedules, and annexes often look like this in production mail-merge outputs.",
    ]
    right_bits = [
        "Column B — technical asides. Process isolation, temp workspaces, and filter pdf:writer_pdf_Export all sit behind a thin Symfony facade.",
        "If columns collapse to a single stream in the PDF, the section properties were lost. If they render but overlap, the gutter/space twips need tuning.",
        "Unicode again in-column: café, niño, Straße, Москва, 東京 — mixed scripts inside narrow measure.",
    ]
    for text in left_bits + right_bits:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            _set_run_font(run, "Liberation Sans", 10)

    # Back to single column before landscape
    single = doc.add_section()
    single.orientation = WD_ORIENT.PORTRAIT
    single.page_width = Cm(21.0)
    single.page_height = Cm(29.7)
    _set_section_columns(single, num=1)

    bridge = doc.add_paragraph()
    run = bridge.add_run("Back to a single-column portrait section before the landscape ledger.")
    _set_run_font(run, "Liberation Sans", 10, italic=True, color=RGBColor(0x64, 0x74, 0x8B))

    # ========== Landscape section ==========
    doc.add_page_break()
    new_section = doc.add_section()
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = Cm(29.7)
    new_section.page_height = Cm(21.0)
    new_section.top_margin = Cm(1.5)
    new_section.bottom_margin = Cm(1.5)
    new_section.left_margin = Cm(1.5)
    new_section.right_margin = Cm(1.5)
    _set_section_columns(new_section, num=1)

    doc.add_heading("8. Landscape section — wide ledger", level=1)
    doc.add_paragraph(
        "This section flips to landscape A4 so LibreOffice must change page geometry mid-document. "
        "Wide tables are a classic fidelity stress for PDF export."
    )

    wide = doc.add_table(rows=1, cols=8)
    wide.style = "Table Grid"
    wide_headers = ["Region", "Q1", "Q2", "Q3", "Q4", "YoY %", "Churn %", "NPS"]
    for i, h in enumerate(wide_headers):
        cell = wide.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, "1E3A5F")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                _set_run_font(run, "Liberation Sans", 9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for region, values in (
        ("Iberia", ("12.4k", "13.1k", "14.0k", "15.2k", "+22.6", "3.1", "62")),
        ("LATAM", ("8.2k", "8.9k", "9.4k", "10.1k", "+23.2", "4.4", "58")),
        ("DACH", ("5.1k", "5.3k", "5.6k", "5.9k", "+15.7", "2.2", "71")),
        ("Nordics", ("2.0k", "2.2k", "2.4k", "2.7k", "+35.0", "2.8", "67")),
        ("RoW", ("1.1k", "1.3k", "1.4k", "1.6k", "+45.5", "5.0", "54")),
    ):
        row = wide.add_row()
        row.cells[0].text = region
        for i, v in enumerate(values, start=1):
            row.cells[i].text = v
            for paragraph in row.cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    run = note.add_run("End of landscape stress section. Returning to portrait is omitted to keep the PDF climax on the wide grid.")
    _set_run_font(run, "Liberation Sans", 10, italic=True, color=RGBColor(0x64, 0x74, 0x8B))

    # Final portrait-ish closing on same landscape is fine; add one more portrait section for closing
    closing = doc.add_section()
    closing.orientation = WD_ORIENT.PORTRAIT
    closing.page_width = Cm(21.0)
    closing.page_height = Cm(29.7)

    doc.add_heading("9. Closing checklist", level=1)
    for item in (
        "Headers/footers with PAGE field rendered",
        "Cover page without repeating header clutter",
        "Two-column section reflows without clipping the gutter",
        "Colored table headers survive export",
        "PNG images appear (not empty frames)",
        "Landscape page(s) present in the PDF",
        "Unicode / emoji glyphs not replaced by boxes (font-dependent)",
    ):
        doc.add_paragraph(item, style="List Bullet")

    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end.paragraph_format.space_before = Pt(24)
    _paragraph_shading(end, "ECFDF5")
    run = end.add_run("✓ If this PDF looks “busy but intact”, the bundle did its job.")
    _set_run_font(run, "Liberation Sans", 12, bold=True, color=RGBColor(0x06, 0x5F, 0x46))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path} ({path.stat().st_size} bytes)")
